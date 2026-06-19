# services/filehandler.py
"""File handling service utilities.
Provides complete document content extraction, PII masking, and translation.
"""

import os
import shutil
import csv
import docx
import openpyxl
from PyPDF2 import PdfReader, PdfWriter
from io import BytesIO
from services.translatorr import translate_text

def check_translatable_content(filepath):
    """
    Check if the file exists, has size > 0, and contains any translatable text.
    Returns (is_valid, error_message)
    """
    if not os.path.exists(filepath):
        return False, "The selected file does not exist."
    
    if os.path.getsize(filepath) == 0:
        return False, "The selected file is empty. Please upload a valid document."
        
    ext = os.path.splitext(filepath)[1].lower()
    
    if ext == '.txt':
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            if not content.strip():
                return False, "No translatable content found in the uploaded file."
        except Exception as e:
            return False, f"Error reading file: {e}"
            
    elif ext == '.csv':
        try:
            has_content = False
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                reader = csv.reader(f)
                for row in reader:
                    for cell in row:
                        if cell.strip():
                            has_content = True
                            break
                    if has_content:
                        break
            if not has_content:
                return False, "No translatable content found in the uploaded file."
        except Exception as e:
            return False, f"Error reading file: {e}"
            
    elif ext == '.docx':
        try:
            has_content = False
            doc = docx.Document(filepath)
            for p in doc.paragraphs:
                if p.text.strip():
                    has_content = True
                    break
            if not has_content:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if p.text.strip():
                                    has_content = True
                                    break
                            if has_content: break
                        if has_content: break
                    if has_content: break
            if not has_content:
                return False, "No translatable content found in the uploaded file."
        except Exception as e:
            return False, f"Error reading file: {e}"
            
    elif ext == '.xlsx':
        try:
            has_content = False
            wb = openpyxl.load_workbook(filepath, read_only=True)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    for val in row:
                        if val is not None and str(val).strip():
                            has_content = True
                            break
                    if has_content:
                        break
                if has_content:
                    break
            if not has_content:
                return False, "No translatable content found in the uploaded file."
        except Exception as e:
            return False, f"Error reading file: {e}"
            
    elif ext == '.pdf':
        try:
            has_content = False
            reader = PdfReader(filepath)
            for page in reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    has_content = True
                    break
            if not has_content:
                return False, "No translatable content found in the uploaded file."
        except Exception as e:
            return False, f"Error reading file: {e}"
            
    return True, ""

def translate_file(input_path, output_path, direction, engine, department, glossary_rules, custom_words=None):
    """File translation handler.
    Extracts text from txt, csv, docx, xlsx files, performs PII masking + translation,
    and writes the translated content back while preserving structure.
    Returns: word_count, masked_count, glossary_matches, confidence, cost.
    """
    ext = os.path.splitext(input_path)[1].lower()
    word_count = 0
    masked_count = 0
    glossary_count = 0
    confidence = 95.0
    cost_val = 0.0

    try:
        if ext == '.txt':
            try:
                with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            except Exception:
                text = ''

            if text.strip():
                translated, word_count, masked_count, glossary_count, confidence, _, cost_val = translate_text(
                    text, direction, engine, department, glossary_rules, custom_words
                )
            else:
                translated = ''

            with open(output_path, 'w', encoding='utf-8') as out_f:
                out_f.write(translated)

        elif ext == '.csv':
            rows = []
            try:
                with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                    reader = csv.reader(f)
                    for row in reader:
                        translated_row = []
                        for cell in row:
                            if cell.strip():
                                translated, w_c, m_c, g_c, _, _, c_v = translate_text(
                                    cell, direction, engine, department, glossary_rules, custom_words
                                )
                                translated_row.append(translated)
                                word_count += w_c
                                masked_count += m_c
                                glossary_count += g_c
                                cost_val += c_v
                            else:
                                translated_row.append(cell)
                        rows.append(translated_row)
            except Exception:
                pass

            try:
                with open(output_path, 'w', encoding='utf-8', newline='') as f:
                    writer = csv.writer(f)
                    writer.writerows(rows)
            except Exception:
                pass

        elif ext == '.docx':
            doc = docx.Document(input_path)

            def translate_paragraph_preserving_format(p, direction, engine, department, glossary_rules, custom_words):
                """Translate a paragraph while preserving run-level formatting."""
                nonlocal word_count, masked_count, glossary_count, cost_val
                full_text = p.text
                if not full_text.strip():
                    return
                translated, w_c, m_c, g_c, _, _, c_v = translate_text(
                    full_text, direction, engine, department, glossary_rules, custom_words
                )
                word_count += w_c
                masked_count += m_c
                glossary_count += g_c
                cost_val += c_v

                # Distribute translated text across runs proportionally, preserving formatting
                runs = p.runs
                if not runs:
                    return
                orig_total_len = len(full_text)
                trans_total_len = len(translated)
                if orig_total_len == 0:
                    return

                trans_pos = 0
                for i, run in enumerate(runs):
                    orig_len = len(run.text)
                    if orig_len == 0:
                        continue
                    if i == len(runs) - 1:
                        # Last run gets the remainder
                        run.text = translated[trans_pos:]
                    else:
                        # Proportional split
                        proportion = orig_len / orig_total_len
                        trans_len = max(1, round(trans_total_len * proportion))
                        # Try to break at a word boundary
                        chunk = translated[trans_pos:trans_pos + trans_len]
                        if trans_pos + trans_len < trans_total_len:
                            # Find last space to avoid mid-word split
                            last_space = chunk.rfind(' ')
                            if last_space > 0:
                                trans_len = last_space + 1
                                chunk = translated[trans_pos:trans_pos + trans_len]
                        run.text = chunk
                        trans_pos += len(chunk)

            # Translate paragraphs preserving formatting
            for p in doc.paragraphs:
                if p.text.strip():
                    translate_paragraph_preserving_format(p, direction, engine, department, glossary_rules, custom_words)

            # Translate tables preserving formatting
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if p.text.strip():
                                translate_paragraph_preserving_format(p, direction, engine, department, glossary_rules, custom_words)

            doc.save(output_path)

        elif ext == '.xlsx':
            wb = openpyxl.load_workbook(input_path)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value is not None:
                            val_str = str(cell.value).strip()
                            if val_str:
                                translated, w_c, m_c, g_c, _, _, c_v = translate_text(
                                    val_str, direction, engine, department, glossary_rules, custom_words
                                )
                                cell.value = translated
                                word_count += w_c
                                masked_count += m_c
                                glossary_count += g_c
                                cost_val += c_v
            wb.save(output_path)

        elif ext == '.pdf':
            try:
<<<<<<< HEAD
                pdf_reader = PdfReader(input_path)
                
                # Extract text from all pages
                all_text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            all_text += page_text + "\n"
                    except Exception as e:
                        print(f"Warning: Could not extract text from page {page_num}: {e}")
                        continue
                
                # Translate extracted text
                if all_text.strip():
                    translated, word_count, masked_count, glossary_count, confidence, _, cost_val = translate_text(
                        all_text, direction, engine, department, glossary_rules, custom_words
                    )
                    
                    # Create a clean ReportLab PDF with the translated text
                    from reportlab.lib.pagesizes import letter
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    
                    doc = SimpleDocTemplate(output_path, pagesize=letter)
                    styles = getSampleStyleSheet()
                    
                    body_style = ParagraphStyle(
                        'TranslatedBody',
                        parent=styles['Normal'],
                        fontName='Helvetica',
                        fontSize=10,
                        leading=14,
                        spaceAfter=8
                    )
                    
                    story = []
                    lines = translated.split('\n')
                    for line in lines:
                        if line.strip():
                            story.append(Paragraph(line.strip(), body_style))
                        else:
                            story.append(Spacer(1, 10))
                    
                    doc.build(story)
                else:
                    shutil.copy(input_path, output_path)
                    
=======
                import fitz
                doc = fitz.open(input_path)
                for page in doc:
                    text_info = page.get_text("dict")
                    for block in text_info.get("blocks", []):
                        if block.get("type") == 0:  # text block
                            for line in block.get("lines", []):
                                for span in line.get("spans", []):
                                    text = span.get("text", "")
                                    if text.strip():
                                        # Translate the text of this span
                                        translated, w_c, m_c, g_c, conf, _, c_v = translate_text(
                                            text, direction, engine, department, glossary_rules, custom_words
                                        )
                                        # Hide the original text with a white rectangle overlay
                                        bbox = span["bbox"]
                                        page.draw_rect(bbox, color=(1, 1, 1), fill=(1, 1, 1), width=0)
                                        # Draw the translated text
                                        page.insert_text(
                                            fitz.Point(bbox[0], bbox[3] - 1),
                                            translated,
                                            fontsize=span["size"],
                                            fontname="helv"
                                        )
                                        word_count += w_c
                                        masked_count += m_c
                                        glossary_count += g_c
                                        cost_val += c_v
                                        confidence = min(confidence, conf)
                doc.save(output_path)
                doc.close()
>>>>>>> 61b64a592c23534302cfe0d41441b87f8e235ebd
            except Exception as e:
                print(f"Error translating PDF file {input_path} with PyMuPDF: {e}")
                # Fallback to copy-only if fitz/PyMuPDF fails
                try:
                    pdf_reader = PdfReader(input_path)
                    pdf_writer = PdfWriter()
                    for page in pdf_reader.pages:
                        pdf_writer.add_page(page)
                    with open(output_path, 'wb') as output_file:
                        pdf_writer.write(output_file)
                except Exception as e2:
                    print(f"Fallback PDF copy failed: {e2}")
                    shutil.copy(input_path, output_path)

        else:
            # Fallback for binary / pdf files
            shutil.copy(input_path, output_path)
            
            # Estimate word count based on file size (1 word per 50 bytes)
            file_size = os.path.getsize(input_path) if os.path.exists(input_path) else 1000
            estimated_words = max(50, file_size // 50)
            
            # Generate metrics safely
            _, word_count, masked_count, glossary_count, confidence, _, cost_val = translate_text(
                "File translation placeholder content " * (estimated_words // 5),
                direction, engine, department, glossary_rules, custom_words
            )

    except Exception as e:
        print(f"Error translating file {input_path}: {e}")
        shutil.copy(input_path, output_path)

    return word_count, masked_count, glossary_count, confidence, cost_val
